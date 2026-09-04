"""Tạo file Word (.docx) từ kết quả phân tích cuộc họp.

Hàm này được chuyển nguyên vẹn từ bizsum-app.py, giữ nguyên mọi logic
và docstring. Chỉ thay đổi: loại bỏ thụt lề cấp 0, đưa import lên đầu file.
"""
import io
import time

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_docx_report(result_json, meeting_title, template_bytes=None):
    """
    Tạo hoặc điền dữ liệu cuộc họp vào file Word (.docx).
    Nếu người dùng truyền template_bytes (.docx), điền các biến placeholder {{TEN_CUOC_HOP}}, {{TOM_TAT}}, {{AM_HUONG}}, v.v.
    Nếu không truyền, tự động tạo mới file Word doanh nghiệp chuẩn mực.
    """
    if isinstance(result_json, list):
        if len(result_json) > 0 and isinstance(result_json[0], dict):
            result_json = result_json[0]
        else:
            result_json = {"summary": str(result_json)}
    elif not isinstance(result_json, dict):
        result_json = {"summary": str(result_json)}

    summary = result_json.get("summary", "Không có tóm tắt.")
    meeting_tone = result_json.get("meeting_tone", "Chuyên nghiệp")
    topics = result_json.get("key_topics", []) or []
    decisions = result_json.get("decisions", []) or []
    actions = result_json.get("action_items", []) or []

    # 1. Trường hợp người dùng tải lên mẫu Word riêng (.docx)
    if template_bytes:
        try:
            doc = docx.Document(io.BytesIO(template_bytes))
            placeholders = {
                "{{TEN_CUOC_HOP}}": meeting_title,
                "{{meeting_title}}": meeting_title,
                "{{TOM_TAT}}": summary,
                "{{summary}}": summary,
                "{{AM_HUONG}}": meeting_tone,
                "{{meeting_tone}}": meeting_tone,
                "{{NGAY_HOP}}": time.strftime("%d/%m/%Y"),
                "{{date}}": time.strftime("%d/%m/%Y"),
            }

            replaced = False

            # Thay thế biến trong các đoạn văn (Paragraphs)
            for p in doc.paragraphs:
                for key, val in placeholders.items():
                    if key in p.text:
                        p.text = p.text.replace(key, str(val))
                        replaced = True

                # Hỗ trợ tự điền vào các mẫu chuẩn tiếng Việt như MauGhiChuCuocHop.docx
                if "Ngày họp:" in p.text and not p.text.strip().endswith(time.strftime("%m/%Y")):
                    p.text = f"Ngày họp: {time.strftime('%d/%m/%Y')}"
                    replaced = True
                elif "Nội dung:" in p.text and len(p.text.strip()) <= 15:
                    content_block = f"Nội dung cuộc họp: {meeting_title}\n\n1. Tóm tắt tổng quan:\n{summary}\n\n2. Âm hưởng cuộc họp: {meeting_tone}\n"
                    if topics:
                        content_block += "\n3. Nội dung thảo luận chi tiết:\n"
                        for idx, topic in enumerate(topics):
                            if isinstance(topic, dict):
                                content_block += f"- {topic.get('topic_name', '')}: {topic.get('discussion_points', '')}\n"
                    if decisions:
                        content_block += "\n4. Các quyết định thông qua:\n"
                        for dec in decisions:
                            content_block += f"- {dec}\n"
                    p.text = content_block
                    replaced = True

            # Thay thế biến trong các Bảng (Tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key, val in placeholders.items():
                                if key in p.text:
                                    p.text = p.text.replace(key, str(val))
                                    replaced = True

            # Nếu có danh sách công việc, tự động bổ sung bảng Action Items vào cuối văn bản mẫu
            if actions:
                doc.add_heading("Kế hoạch hành động & Phân công công việc (Action Items)", level=2)
                table = doc.add_table(rows=1, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"

                hdr_cells = table.rows[0].cells
                hdr_titles = ["Nhiệm vụ cần làm", "Người chịu trách nhiệm", "Hạn chót hoàn thành"]
                for i, title in enumerate(hdr_titles):
                    hdr_cells[i].text = title
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(15, 23, 42)

                for act in actions:
                    if isinstance(act, dict):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(act.get("task", "Chưa rõ"))
                        row_cells[1].text = str(act.get("assignee", "Chưa chỉ định"))
                        row_cells[2].text = str(act.get("deadline", "Chưa rõ"))

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception:
            pass

    # 2. Trường hợp tạo mới file Word doanh nghiệp chuẩn mực
    doc = docx.Document()

    # Tiêu đề chính
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run(f"BIÊN BẢN HỌP: {meeting_title.upper()}")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    # Ngày lập biên bản
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = date_p.add_run(f"Thời gian lập: {time.strftime('%d/%m/%Y %H:%M')}")
    run_date.font.italic = True
    run_date.font.size = Pt(10)
    run_date.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # Phần 1: Tóm tắt tổng quan
    doc.add_heading("1. Tóm tắt tổng quan cuộc họp", level=1)
    p1 = doc.add_paragraph(summary)
    p1.paragraph_format.line_spacing = 1.25

    # Phần 2: Trạng thái & Âm hưởng
    doc.add_heading("2. Trạng thái & Âm hưởng cuộc họp", level=1)
    doc.add_paragraph(f"- Âm hưởng chung: {meeting_tone}")
    doc.add_paragraph(f"- Tổng số nhiệm vụ phân công: {len(actions)} hạng mục công việc")

    # Phần 3: Các chủ đề thảo luận chính
    doc.add_heading("3. Nội dung thảo luận chi tiết", level=1)
    if topics:
        for idx, topic in enumerate(topics):
            if isinstance(topic, dict):
                t_name = topic.get("topic_name", f"Chủ đề {idx+1}")
                t_points = topic.get("discussion_points", "")
                p_topic = doc.add_paragraph()
                r_topic = p_topic.add_run(f"3.{idx+1} {t_name}")
                r_topic.bold = True
                doc.add_paragraph(t_points)
    else:
        doc.add_paragraph("Không phát hiện chủ đề thảo luận cụ thể.")

    # Phần 4: Các quyết định thông qua
    doc.add_heading("4. Các quyết định đã được thông qua", level=1)
    if decisions:
        for dec in decisions:
            doc.add_paragraph(f"• {dec}", style="List Bullet")
    else:
        doc.add_paragraph("Chưa ghi nhận quyết định thống nhất chính thức nào.")

    # Phần 5: Phân công công việc (Bảng Action Items)
    doc.add_heading("5. Kế hoạch hành động & Phân công công việc (Action Items)", level=1)
    if actions:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_titles = ["Nhiệm vụ cần làm", "Người chịu trách nhiệm", "Hạn chót hoàn thành"]
        for i, title in enumerate(hdr_titles):
            hdr_cells[i].text = title
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(15, 23, 42)

        for act in actions:
            if isinstance(act, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = str(act.get("task", "Chưa rõ"))
                row_cells[1].text = str(act.get("assignee", "Chưa chỉ định"))
                row_cells[2].text = str(act.get("deadline", "Chưa rõ"))
    else:
        doc.add_paragraph("Không có công việc cụ thể được phân công.")

    # Chữ ký xác nhận
    doc.add_paragraph()
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig_p.add_run("Người lập biên bản\n(Ký và ghi rõ họ tên)\n\n\n_______________________")
    sig_run.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
