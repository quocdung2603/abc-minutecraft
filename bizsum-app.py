import streamlit as st
import google.generativeai as genai
import json
import os
import tempfile
import time
import io
import pandas as pd
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image

# Cấu hình trang Streamlit chuẩn Doanh nghiệp
st.set_page_config(
    page_title="MinuteCraft - Nền Tảng Tự Động Hóa Biên Bản Họp SME",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Giao diện CSS chuẩn Production (Tự động thích ứng Light Mode & Dark Mode)
st.markdown("""
<style>
    /* Dynamic Theme CSS Variables */
    :root {
        --bg-main: #F8FAFC;
        --card-bg: #FFFFFF;
        --card-border: #E2E8F0;
        --text-primary: #0F172A;
        --text-secondary: #475569;
        --kpi-label: #64748B;
        --kpi-border: #E2E8F0;
        --accent-blue: #2563EB;
        --svg-stroke: #2563EB;
    }

    /* Dark Mode Auto-Adaptation */
    @media (prefers-color-scheme: dark) {
        :root {
            --bg-main: #0F172A;
            --card-bg: #1E293B;
            --card-border: #334155;
            --text-primary: #F8FAFC;
            --text-secondary: #94A3B8;
            --kpi-label: #94A3B8;
            --kpi-border: #334155;
            --accent-blue: #3B82F6;
            --svg-stroke: #60A5FA;
        }
    }

    /* Streamlit Dark Theme Class Overrides */
    [data-theme="dark"] {
        --bg-main: #0F172A;
        --card-bg: #1E293B;
        --card-border: #334155;
        --text-primary: #F8FAFC;
        --text-secondary: #94A3B8;
        --kpi-label: #94A3B8;
        --kpi-border: #334155;
        --accent-blue: #3B82F6;
        --svg-stroke: #60A5FA;
    }
    
    /* Executive Main Title */
    .main-title {
        font-size: 2.2rem;
        font-weight: 800;
        color: var(--text-primary);
        letter-spacing: -0.02em;
        margin-bottom: 0.25rem;
    }
    .subtitle {
        font-size: 1.05rem;
        color: var(--text-secondary);
        margin-bottom: 1.75rem;
    }
    
    /* Production Cards */
    .prod-card {
        background: var(--card-bg);
        border: 1px solid var(--card-border);
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1.25rem;
        box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.05);
    }
    .prod-card-title {
        font-size: 1.1rem;
        font-weight: 700;
        color: var(--text-primary);
        margin-bottom: 0.75rem;
        display: flex;
        align-items: center;
        gap: 8px;
    }
    
    /* Custom KPI Metric Grid */
    .kpi-container {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1rem;
        margin-bottom: 1.5rem;
    }
    .kpi-box {
        background: var(--card-bg);
        border: 1px solid var(--kpi-border);
        border-radius: 10px;
        padding: 1rem 1.25rem;
        box-shadow: 0 1px 2px 0 rgba(0, 0, 0, 0.03);
    }
    .kpi-box-label {
        font-size: 0.75rem;
        font-weight: 600;
        color: var(--kpi-label);
        text-transform: uppercase;
        letter-spacing: 0.05em;
        margin-bottom: 0.25rem;
    }
    .kpi-box-value {
        font-size: 1.15rem;
        font-weight: 700;
        color: var(--text-primary);
    }
    
    /* Primary CTA Button */
    .stButton > button {
        background-color: var(--accent-blue) !important;
        color: white !important;
        font-weight: 700 !important;
        font-size: 1.05rem !important;
        padding: 0.75rem 2rem !important;
        border-radius: 8px !important;
        border: none !important;
        box-shadow: 0 4px 6px -1px rgba(37, 99, 235, 0.2) !important;
        transition: all 0.2s ease !important;
    }
    .stButton > button:hover {
        opacity: 0.9 !important;
        box-shadow: 0 6px 12px -2px rgba(37, 99, 235, 0.3) !important;
    }
</style>
""", unsafe_allow_html=True)


def generate_docx_report(result_json, meeting_title, template_bytes=None):
    """
    Tạo hoặc điền dữ liệu cuộc họp vào file Word (.docx).
    Nếu người dùng truyền template_bytes (.docx), điền các biến placeholder {{TEN_CUOC_HOP}}, {{TOM_TAT}}, {{AM_HUONG}}, v.v.
    Nếu không truyền, tự động tạo mới file Word doanh nghiệp chuẩn mực.
    """
    if isinstance(result_json, list):
        if len(result_json) > 0 and isinstance(result_json[0], dict):
            result_json = result_json[0]
        else:
            result_json = {"summary": str(result_json)}
    elif not isinstance(result_json, dict):
        result_json = {"summary": str(result_json)}

    summary = result_json.get("summary", "Không có tóm tắt.")
    meeting_tone = result_json.get("meeting_tone", "Chuyên nghiệp")
    topics = result_json.get("key_topics", []) or []
    decisions = result_json.get("decisions", []) or []
    actions = result_json.get("action_items", []) or []

    # 1. Trường hợp người dùng tải lên mẫu Word riêng (.docx)
    if template_bytes:
        try:
            doc = docx.Document(io.BytesIO(template_bytes))
            placeholders = {
                "{{TEN_CUOC_HOP}}": meeting_title,
                "{{meeting_title}}": meeting_title,
                "{{TOM_TAT}}": summary,
                "{{summary}}": summary,
                "{{AM_HUONG}}": meeting_tone,
                "{{meeting_tone}}": meeting_tone,
                "{{NGAY_HOP}}": time.strftime("%d/%m/%Y"),
                "{{date}}": time.strftime("%d/%m/%Y"),
            }
            
            replaced = False

            # Thay thế biến trong các đoạn văn (Paragraphs)
            for p in doc.paragraphs:
                for key, val in placeholders.items():
                    if key in p.text:
                        p.text = p.text.replace(key, str(val))
                        replaced = True

                # Hỗ trợ tự điền vào các mẫu chuẩn tiếng Việt như MauGhiChuCuocHop.docx
                if "Ngày họp:" in p.text and not p.text.strip().endswith(time.strftime("%m/%Y")):
                    p.text = f"Ngày họp: {time.strftime('%d/%m/%Y')}"
                    replaced = True
                elif "Nội dung:" in p.text and len(p.text.strip()) <= 15:
                    content_block = f"Nội dung cuộc họp: {meeting_title}\n\n1. Tóm tắt tổng quan:\n{summary}\n\n2. Âm hưởng cuộc họp: {meeting_tone}\n"
                    if topics:
                        content_block += "\n3. Nội dung thảo luận chi tiết:\n"
                        for idx, topic in enumerate(topics):
                            if isinstance(topic, dict):
                                content_block += f"- {topic.get('topic_name', '')}: {topic.get('discussion_points', '')}\n"
                    if decisions:
                        content_block += "\n4. Các quyết định thông qua:\n"
                        for dec in decisions:
                            content_block += f"- {dec}\n"
                    p.text = content_block
                    replaced = True

            # Thay thế biến trong các Bảng (Tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key, val in placeholders.items():
                                if key in p.text:
                                    p.text = p.text.replace(key, str(val))
                                    replaced = True

            # Nếu có danh sách công việc, tự động bổ sung bảng Action Items vào cuối văn bản mẫu
            if actions:
                doc.add_heading("Kế hoạch hành động & Phân công công việc (Action Items)", level=2)
                table = doc.add_table(rows=1, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"

                hdr_cells = table.rows[0].cells
                hdr_titles = ["Nhiệm vụ cần làm", "Người chịu trách nhiệm", "Hạn chót hoàn thành"]
                for i, title in enumerate(hdr_titles):
                    hdr_cells[i].text = title
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(15, 23, 42)

                for act in actions:
                    if isinstance(act, dict):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(act.get("task", "Chưa rõ"))
                        row_cells[1].text = str(act.get("assignee", "Chưa chỉ định"))
                        row_cells[2].text = str(act.get("deadline", "Chưa rõ"))
                                    
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception:
            pass

    # 2. Trường hợp tạo mới file Word doanh nghiệp chuẩn mực
    doc = docx.Document()

    # Tiêu đề chính
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run(f"BIÊN BẢN HỌP: {meeting_title.upper()}")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    # Ngày lập biên bản
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = date_p.add_run(f"Thời gian lập: {time.strftime('%d/%m/%Y %H:%M')}")
    run_date.font.italic = True
    run_date.font.size = Pt(10)
    run_date.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # Phần 1: Tóm tắt tổng quan
    doc.add_heading("1. Tóm tắt tổng quan cuộc họp", level=1)
    p1 = doc.add_paragraph(summary)
    p1.paragraph_format.line_spacing = 1.25

    # Phần 2: Trạng thái & Âm hưởng
    doc.add_heading("2. Trạng thái & Âm hưởng cuộc họp", level=1)
    doc.add_paragraph(f"- Âm hưởng chung: {meeting_tone}")
    doc.add_paragraph(f"- Tổng số nhiệm vụ phân công: {len(actions)} hạng mục công việc")

    # Phần 3: Các chủ đề thảo luận chính
    doc.add_heading("3. Nội dung thảo luận chi tiết", level=1)
    if topics:
        for idx, topic in enumerate(topics):
            if isinstance(topic, dict):
                t_name = topic.get("topic_name", f"Chủ đề {idx+1}")
                t_points = topic.get("discussion_points", "")
                p_topic = doc.add_paragraph()
                r_topic = p_topic.add_run(f"3.{idx+1} {t_name}")
                r_topic.bold = True
                doc.add_paragraph(t_points)
    else:
        doc.add_paragraph("Không phát hiện chủ đề thảo luận cụ thể.")

    # Phần 4: Các quyết định thông qua
    doc.add_heading("4. Các quyết định đã được thông qua", level=1)
    if decisions:
        for dec in decisions:
            doc.add_paragraph(f"• {dec}", style="List Bullet")
    else:
        doc.add_paragraph("Chưa ghi nhận quyết định thống nhất chính thức nào.")

    # Phần 5: Phân công công việc (Bảng Action Items)
    doc.add_heading("5. Kế hoạch hành động & Phân công công việc (Action Items)", level=1)
    if actions:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_titles = ["Nhiệm vụ cần làm", "Người chịu trách nhiệm", "Hạn chót hoàn thành"]
        for i, title in enumerate(hdr_titles):
            hdr_cells[i].text = title
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(15, 23, 42)

        for act in actions:
            if isinstance(act, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = str(act.get("task", "Chưa rõ"))
                row_cells[1].text = str(act.get("assignee", "Chưa chỉ định"))
                row_cells[2].text = str(act.get("deadline", "Chưa rõ"))
    else:
        doc.add_paragraph("Không có công việc cụ thể được phân công.")

    # Chữ ký xác nhận
    doc.add_paragraph()
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig_p.add_run("Người lập biên bản\n(Ký và ghi rõ họ tên)\n\n\n_______________________")
    sig_run.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------
# TỰ ĐỘNG LƯU & CACHE API KEY (API KEY PERSISTENCE)
# ---------------------------------------------------------
KEY_FILE = ".api_key_cache"

if "gemini_api_key" not in st.session_state:
    if os.path.exists(KEY_FILE):
        try:
            with open(KEY_FILE, "r", encoding="utf-8") as f:
                st.session_state["gemini_api_key"] = f.read().strip()
        except Exception:
            st.session_state["gemini_api_key"] = ""
    else:
        st.session_state["gemini_api_key"] = ""

saved_key = st.session_state.get("gemini_api_key", "")

# Sidebar Cấu Hình
st.sidebar.markdown("""
<div style="display: flex; align-items: center; gap: 8px; margin-bottom: 1rem;">
    <svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="var(--svg-stroke)" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
        <circle cx="12" cy="12" r="3"/>
        <path d="M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 0 1 0 2.83 2 2 0 0 1-2.83 0l-.06-.06a1.65 1.65 0 0 0-1.82-.33 1.65 1.65 0 0 0-1 1.51V21a2 2 0 0 1-2 2 2 2 0 0 1-2-2v-.09A1.65 1.65 0 0 0 9 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 0 1-2.83 0 2 2 0 0 1 0-2.83l.06-.06a1.65 1.65 0 0 0 .33-1.82 1.65 1.65 0 0 0-1.51-1H3a2 2 0 0 1-2-2 2 2 0 0 1 2-2h.09A1.65 1.65 0 0 0 4.6 9a1.65 1.65 0 0 0-.33-1.82l-.06-.06a2 2 0 0 1 0-2.83 2 2 0 0 1 2.83 0l.06.06a1.65 1.65 0 0 0 1.82.33H9a1.65 1.65 0 0 0 1-1.51V3a2 2 0 0 1 2-2 2 2 0 0 1 2 2v.09a1.65 1.65 0 0 0 1 1.51 1.65 1.65 0 0 0 1.82-.33l.06-.06a2 2 0 0 1 2.83 0 2 2 0 0 1 0 2.83l-.06.06a1.65 1.65 0 0 0-.33 1.82V9a1.65 1.65 0 0 0 1.51 1H21a2 2 0 0 1 2 2 2 2 0 0 1-2 2h-.09a1.65 1.65 0 0 0-1.51 1z"/>
    </svg>
    <h3 style="margin: 0; color: var(--text-primary); font-weight: 700; font-size: 1.2rem;">Cấu Hình Hệ Thống</h3>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("### Gemini API Key")
api_key_input = st.sidebar.text_input(
    "Khóa API của bạn", 
    value=saved_key, 
    type="password",
    help="Khóa API được tự động lưu cho các lần truy cập sau."
)

if api_key_input:
    api_key_input = api_key_input.strip()
    if api_key_input != saved_key:
        st.session_state["gemini_api_key"] = api_key_input
        try:
            with open(KEY_FILE, "w", encoding="utf-8") as f:
                f.write(api_key_input)
        except Exception:
            pass

api_key = st.session_state.get("gemini_api_key", "")

if api_key:
    st.sidebar.success("Đã kết nối API Key")
    if st.sidebar.button("Xóa API Key đã lưu"):
        st.session_state["gemini_api_key"] = ""
        if os.path.exists(KEY_FILE):
            try:
                os.remove(KEY_FILE)
            except Exception:
                pass
        st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("### Thông số Mô hình AI")
model_choice = st.sidebar.selectbox(
    "Mô hình AI", 
    ["gemini-3.6-flash", "gemini-2.5-flash", "gemini-flash-latest", "gemini-1.5-flash", "gemini-1.5-pro"], 
    index=0, 
    help="gemini-3.6-flash là mô hình mới nhất được khuyến nghị."
)
temperature = st.sidebar.slider(
    "Độ sáng tạo (Temperature)", 
    min_value=0.0, max_value=1.0, value=0.1, step=0.1,
    help="Mức 0.1 đảm bảo biên bản họp chính xác, bám sát thực tế cuộc họp."
)

st.sidebar.markdown("---")
st.sidebar.markdown("### Hướng dẫn nhanh")
st.sidebar.info("""
1. Nhập API Key một lần duy nhất (Hệ thống tự ghi nhớ).
2. Tải lên tệp âm thanh `.mp3`, `.wav`, hoặc `.m4a` (hỗ trợ tối đa 100MB).
3. Bấm 'Bắt đầu Phân tích' để tự động lập biên bản họp.
""")

# ---------------------------------------------------------
# HEADER CHÍNH ỨNG DỤNG
# ---------------------------------------------------------
st.markdown("""
<div style="display: flex; align-items: center; gap: 12px; margin-bottom: 0.25rem;">
    <svg width="34" height="34" viewBox="0 0 24 24" fill="none" stroke="var(--svg-stroke)" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round">
        <path d="M14.5 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V7.5L14.5 2z"/>
        <polyline points="14 2 14 8 20 8"/>
        <path d="M16 13H8"/>
        <path d="M16 17H8"/>
        <path d="M10 9H8"/>
    </svg>
    <span class="main-title">MinuteCraft</span>
</div>
""", unsafe_allow_html=True)
st.markdown("<div class='subtitle'>Nền tảng tự động hóa biên bản ghi chép cuộc họp cho doanh nghiệp SME</div>", unsafe_allow_html=True)

if not api_key:
    st.warning("Vui lòng nhập Gemini API Key ở thanh bên trái (Sidebar) để bắt đầu!")
    st.stop()

genai.configure(api_key=api_key)

# ---------------------------------------------------------
# KHU VỰC TẢI FILE & CẤU HÌNH NGHIỆP VỤ HỌP
# ---------------------------------------------------------
st.markdown("""
<div style="display: flex; align-items: center; gap: 10px; margin-top: 0.5rem; margin-bottom: 0.75rem;">
    <svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="var(--svg-stroke)" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
        <path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/>
        <polyline points="17 8 12 3 7 8"/>
        <line x1="12" y1="3" x2="12" y2="15"/>
    </svg>
    <h3 style="margin: 0; color: var(--text-primary); font-weight: 700; font-size: 1.25rem;">Tải Lên File Ghi Âm Cuộc Họp</h3>
</div>
""", unsafe_allow_html=True)

uploaded_file = st.file_uploader("Chọn hoặc kéo thả tệp âm thanh cuộc họp tại đây", type=["mp3", "wav", "m4a", "ogg"])

if uploaded_file is not None:
    col_audio, col_config = st.columns([1, 1.2], gap="medium")
    
    with col_audio:
        st.markdown("""
        <div class="prod-card">
            <div class="prod-card-title">
                <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="var(--svg-stroke)" stroke-width="2"><path d="M9 18V5l12-2v13"/><circle cx="6" cy="18" r="3"/><circle cx="18" cy="16" r="3"/></svg>
                Thông tin tệp ghi âm
            </div>
        """, unsafe_allow_html=True)
        file_details = {
            "Tên tệp": uploaded_file.name,
            "Kích thước": f"{uploaded_file.size / (1024*1024):.2f} MB",
            "Định dạng": uploaded_file.type or "Audio"
        }
        st.json(file_details)
        st.audio(uploaded_file.getvalue(), format=uploaded_file.type or "audio/wav")
        st.markdown("</div>", unsafe_allow_html=True)
        
    with col_config:
        st.markdown("""
        <div class="prod-card">
            <div class="prod-card-title">
                <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="var(--svg-stroke)" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>
                Yêu cầu trọng tâm phân tích
            </div>
        """, unsafe_allow_html=True)
        focus_area = st.selectbox("Trọng tâm phân tích", [
            "Tóm tắt toàn diện (Tổng quan, Quyết định, Phân công công việc)",
            "Chỉ tập trung vào Danh sách Việc cần làm (Action Items)",
            "Tập trung vào các Ý kiến đóng góp và Tranh luận (Brainstorming Details)",
            "Trích xuất nhanh các Mốc thời gian & Deadline quan trọng"
        ])
        
        # Thêm lựa chọn Mẫu Word (.docx) doanh nghiệp
        template_file = st.file_uploader(
            "Tải lên Mẫu biên bản Word (.docx) của công ty bạn (Không bắt buộc)", 
            type=["docx"],
            help="Hệ thống sẽ tự động điền các biến {{TEN_CUOC_HOP}}, {{TOM_TAT}}, {{AM_HUONG}}, {{NGAY_HOP}} vào mẫu Word của bạn."
        )

        custom_instructions = st.text_area(
            "Yêu cầu bổ sung (Không bắt buộc):", 
            placeholder="Ví dụ: Định dạng ngày tháng dạng DD/MM/YYYY; Cuộc họp sử dụng nhiều thuật ngữ chuyên ngành..."
        )
        st.markdown("</div>", unsafe_allow_html=True)

    # ---------------------------------------------------------
    # NÚT KÍCH HOẠT PHÂN TÍCH
    # ---------------------------------------------------------
    if st.button("Bắt đầu Phân tích Cuộc họp", use_container_width=True):
        with st.spinner("Hệ thống đang xử lý âm thanh và truyền đến Gemini Cloud API..."):
            try:
                file_bytes = uploaded_file.getvalue()
                mime_type = getattr(uploaded_file, "type", None) or "audio/mp3"
                if mime_type not in ["audio/mp3", "audio/wav", "audio/m4a", "audio/ogg", "audio/x-m4a"]:
                    ext = os.path.splitext(uploaded_file.name)[1].lower()
                    if ext == ".wav":
                        mime_type = "audio/wav"
                    elif ext in [".m4a", ".aac"]:
                        mime_type = "audio/m4a"
                    elif ext == ".ogg":
                        mime_type = "audio/ogg"
                    else:
                        mime_type = "audio/mp3"

                audio_payload = None
                gemini_file = None
                tmp_file_path = None

                if len(file_bytes) <= 100 * 1024 * 1024:
                    audio_payload = {"mime_type": mime_type, "data": file_bytes}
                else:
                    try:
                        file_ext = os.path.splitext(uploaded_file.name)[1]
                        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                            tmp_file.write(file_bytes)
                            tmp_file_path = tmp_file.name

                        st.info("Đang tải tệp âm thanh lớn lên Gemini Cloud File API...")
                        gemini_file = genai.upload_file(path=tmp_file_path, mime_type=mime_type)
                        while gemini_file.state.name == "PROCESSING":
                            time.sleep(2)
                            gemini_file = genai.get_file(gemini_file.name)
                        
                        if gemini_file.state.name == "FAILED":
                            st.error("Quá trình tải lên tệp tin thất bại!")
                            st.stop()
                        audio_payload = gemini_file
                    except Exception as e_upload:
                        audio_payload = {"mime_type": mime_type, "data": file_bytes}

                st.success("Tải tệp thành công! Đang tiến hành phân tích nội dung cuộc trò chuyện...")

                system_instruction = """
                Bạn là một thư ký cuộc họp ảo chuyên nghiệp hoạt động trong môi trường doanh nghiệp vừa và nhỏ (SME) tại Việt Nam.
                Nhiệm vụ của bạn là lắng nghe tệp ghi âm cuộc họp, thấu hiểu ngữ cảnh giao tiếp bằng tiếng Việt (bao gồm cả các từ ngữ viết tắt, từ lóng công nghệ hoặc thuật ngữ kinh doanh), và tự động lập một biên bản họp cực kỳ chuyên nghiệp, ngắn gọn, súc tích và chính xác.
                
                Bạn PHẢI phân tích và trích xuất thông tin theo cấu trúc JSON duy nhất (không bọc trong bất kỳ đoạn text giải thích nào khác bên ngoài JSON):
                {
                  "meeting_title": "Tên cuộc họp (AI tự suy luận dựa trên nội dung)",
                  "summary": "Tóm tắt tổng quan về mục tiêu và nội dung chính của cuộc họp (từ 3 đến 5 dòng)",
                  "key_topics": [
                    {
                      "topic_name": "Tên chủ đề trao đổi 1",
                      "discussion_points": "Tóm tắt các ý kiến, tranh luận của các thành viên về chủ đề này"
                    }
                  ],
                  "decisions": [
                    "Các quyết định quan trọng đã được thống nhất 1",
                    "Các quyết định quan trọng đã được thống nhất 2"
                  ],
                  "action_items": [
                    {
                      "task": "Nhiệm vụ cụ thể cần thực hiện",
                      "assignee": "Người chịu trách nhiệm (nếu không rõ, ghi 'Chưa chỉ định')",
                      "deadline": "Hạn chót hoàn thành (nếu không có, ghi 'Chưa rõ')"
                    }
                  ],
                  "meeting_tone": "Đánh giá không khí cuộc họp (ví dụ: Chuyên nghiệp, Khẩn trương, Căng thẳng, Cởi mở...)"
                }
                """

                prompt = f"Hãy nghe file âm thanh này và lập biên bản họp chi tiết theo cấu trúc JSON được hướng dẫn."
                if focus_area == "Chỉ tập trung vào Danh sách Việc cần làm (Action Items)":
                    prompt += " Đặc biệt tập trung bóc tách thật chi tiết các nhiệm vụ (Action Items) được giao."
                elif focus_area == "Tập trung vào các Ý kiến đóng góp và Tranh luận (Brainstorming Details)":
                    prompt += " Đặc biệt tập trung ghi nhận chi tiết các luồng tranh biện, ý kiến đóng góp khác nhau của các thành viên."
                elif focus_area == "Trích xuất nhanh các Mốc thời gian & Deadline quan trọng":
                    prompt += " Đặc biệt chú ý bóc tách rõ các cột mốc thời gian và thời hạn hoàn thành công việc."

                if custom_instructions:
                    prompt += f" Lưu ý thêm các yêu cầu đặc biệt sau từ người dùng: {custom_instructions}"

                candidate_models = ["gemini-3.6-flash", model_choice, "gemini-2.5-flash", "gemini-flash-latest", "gemini-1.5-flash"]
                seen = set()
                candidate_models = [m for m in candidate_models if not (m in seen or seen.add(m))]

                response = None
                last_err = None
                for m_name in candidate_models:
                    try:
                        model = genai.GenerativeModel(
                            model_name=m_name,
                            generation_config={
                                "temperature": temperature,
                                "response_mime_type": "application/json"
                            },
                            system_instruction=system_instruction
                        )
                        response = model.generate_content([audio_payload, prompt])
                        break
                    except Exception as e_model:
                        last_err = e_model
                        err_str = str(e_model)
                        if "404" in err_str or "is no longer available" in err_str or "not found" in err_str:
                            continue
                        else:
                            raise e_model
                
                if response is None and last_err is not None:
                    raise last_err

                if tmp_file_path and os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                if gemini_file:
                    try:
                        genai.delete_file(gemini_file.name)
                    except Exception:
                        pass

                raw_text = response.text.strip()
                if raw_text.startswith("```json"):
                    raw_text = raw_text[7:]
                if raw_text.startswith("```"):
                    raw_text = raw_text[3:]
                if raw_text.endswith("```"):
                    raw_text = raw_text[:-3]
                raw_text = raw_text.strip()

                result_json = json.loads(raw_text)
                
                # Chuẩn hóa dữ liệu JSON: Đảm bảo result_json luôn là dict
                if isinstance(result_json, list):
                    if len(result_json) > 0 and isinstance(result_json[0], dict):
                        result_json = result_json[0]
                    else:
                        result_json = {"summary": str(result_json)}
                elif not isinstance(result_json, dict):
                    result_json = {"summary": str(result_json)}
                
                # ---------------------------------------------------------
                # HIỂN THỊ KẾT QUẢ PHÂN TÍCH CHUẨN PRODUCTION
                # ---------------------------------------------------------
                st.balloons()
                st.success("Đã hoàn thành biên bản họp tự động thành công!")
                
                meeting_title = result_json.get("meeting_title", "Cuộc họp không có tiêu đề") or "Cuộc họp không có tiêu đề"
                actions = result_json.get("action_items", []) or []
                
                st.markdown(f"## Biên bản họp: {meeting_title}")
                
                # Custom Production KPI Grid
                st.markdown(f"""
                <div class="kpi-container">
                    <div class="kpi-box">
                        <div class="kpi-box-label">Không khí cuộc họp</div>
                        <div class="kpi-box-value">{result_json.get('meeting_tone', 'Chuyên nghiệp')}</div>
                    </div>
                    <div class="kpi-box">
                        <div class="kpi-box-label">Tổng số nhiệm vụ</div>
                        <div class="kpi-box-value">{len(actions)} Công việc</div>
                    </div>
                    <div class="kpi-box">
                        <div class="kpi-box-label">Dung lượng tệp</div>
                        <div class="kpi-box-value">{uploaded_file.size / (1024*1024):.1f} MB</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # Khung Tóm Tắt Tổng Quan
                st.markdown("""
                <div class="prod-card">
                    <div class="prod-card-title">
                        <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="var(--svg-stroke)" stroke-width="2"><line x1="8" y1="6" x2="21" y2="6"/><line x1="8" y1="12" x2="21" y2="12"/><line x1="8" y1="18" x2="21" y2="18"/><line x1="3" y1="6" x2="3.01" y2="6"/><line x1="3" y1="12" x2="3.01" y2="12"/><line x1="3" y1="18" x2="3.01" y2="18"/></svg>
                        Tóm tắt tổng quan cuộc họp
                    </div>
                """, unsafe_allow_html=True)
                st.write(result_json.get("summary", "Không có tóm tắt."))
                st.markdown("</div>", unsafe_allow_html=True)

                # Tabs Chi Tiết
                tab1, tab2, tab3 = st.tabs(["Nội dung trao đổi chi tiết", "Quyết định thống nhất", "Kế hoạch hành động (Action Items)"])
                
                topics = result_json.get("key_topics", []) or []
                with tab1:
                    st.markdown("### Các chủ đề thảo luận chính")
                    if topics:
                        for idx, topic in enumerate(topics):
                            t_name = topic.get("topic_name", f"Chủ đề {idx+1}") if isinstance(topic, dict) else f"Chủ đề {idx+1}"
                            t_points = topic.get("discussion_points", "") if isinstance(topic, dict) else str(topic)
                            with st.expander(f"Chủ đề {idx+1}: {t_name}", expanded=True):
                                st.write(t_points)
                    else:
                        st.write("Không phát hiện chủ đề thảo luận cụ thể.")

                decisions = result_json.get("decisions", []) or []
                with tab2:
                    st.markdown("### Các quyết định đã được thông qua")
                    if decisions:
                        for dec in decisions:
                            st.markdown(f"- **{dec}**")
                    else:
                        st.info("Cuộc họp chưa ghi nhận quyết định thống nhất chính thức nào.")

                with tab3:
                    st.markdown("### Phân công công việc & Deadline chi tiết")
                    if actions:
                        formatted_actions = []
                        for act in actions:
                            if isinstance(act, dict):
                                formatted_actions.append({
                                    "Nhiệm vụ cần làm": act.get("task", "Chưa rõ"),
                                    "Người chịu trách nhiệm": act.get("assignee", "Chưa chỉ định"),
                                    "Hạn chót hoàn thành": act.get("deadline", "Chưa rõ")
                                })
                        df = pd.DataFrame(formatted_actions)
                        st.dataframe(df, use_container_width=True)
                    else:
                        st.info("Không phát hiện nhiệm vụ được giao cụ thể trong cuộc họp.")

                # Xuất Báo Cáo Song Song (Word .docx & Markdown .md)
                st.markdown("---")
                st.subheader("Tải Xuống Biên Bản Họp (Đa định dạng)")
                
                report_md = f"# BIÊN BẢN HỌP TỰ ĐỘNG: {meeting_title}\n\n"
                report_md += f"## 1. Tóm tắt tổng quan\n{result_json.get('summary', 'Không có tóm tắt.')}\n\n"
                report_md += f"## 2. Trạng thái cuộc họp\n- Không khí: {result_json.get('meeting_tone', 'Chuyên nghiệp')}\n\n"
                report_md += "## 3. Nội dung thảo luận chi tiết\n"
                for topic in topics:
                    if isinstance(topic, dict):
                        report_md += f"### - {topic.get('topic_name', '')}\n{topic.get('discussion_points', '')}\n\n"
                report_md += "## 4. Các quyết định thông qua\n"
                for dec in decisions:
                    report_md += f"- [x] {dec}\n"
                report_md += "\n## 5. Phân công công việc (Action Items)\n"
                for act in actions:
                    if isinstance(act, dict):
                        report_md += f"- **Công việc**: {act.get('task', '')} | **Người làm**: {act.get('assignee', '')} | **Hạn chót**: {act.get('deadline', '')}\n"
                
                sanitized_title = "".join([c if c.isalnum() or c in (" ", "_", "-") else "" for c in meeting_title]).strip().replace(" ", "_")
                if not sanitized_title:
                    sanitized_title = "Bien_ban_hop"

                # Tạo dữ liệu Word (.docx)
                template_bytes = template_file.getvalue() if template_file is not None else None
                docx_bytes = generate_docx_report(result_json, meeting_title, template_bytes)

                col_dl1, col_dl2 = st.columns(2, gap="medium")
                with col_dl1:
                    st.download_button(
                        label="📄 Tải biên bản họp dạng Word (.docx)",
                        data=docx_bytes,
                        file_name=f"Bien_ban_hop_{sanitized_title}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with col_dl2:
                    st.download_button(
                        label="📝 Tải biên bản họp dạng Markdown (.md)",
                        data=report_md.encode("utf-8"),
                        file_name=f"Bien_ban_hop_{sanitized_title}.md",
                        mime="text/markdown",
                        use_container_width=True
                    )

            except Exception as e:
                err_msg = str(e)
                if "429" in err_msg or "Quota exceeded" in err_msg or "ResourceExhausted" in err_msg:
                    st.warning("**Vượt quá giới hạn (Rate Limit / Quota Exceeded)!** Bạn đang gửi quá nhiều yêu cầu liên tục (Free tier giới hạn 5-15 request/phút). Vui lòng đợi khoảng 15-30 giây rồi bấm thử lại!")
                elif "API_KEY_INVALID" in err_msg or "API key not valid" in err_msg:
                    st.error("**Lỗi API Key không hợp lệ!** Gemini API Key bạn nhập không chính xác hoặc đã bị khóa/vô hiệu hóa.")
                    st.warning("**Cách khắc phục:**\n1. Truy cập [Google AI Studio (aistudio.google.com)](https://aistudio.google.com/)\n2. Bấm **Get API key** và tạo khóa mới.\n3. Dán khóa mới vào thanh Sidebar bên trái và bấm chạy lại.")
                else:
                    st.error(f"Đã xảy ra lỗi trong quá trình xử lý: {err_msg}")
                    st.info("Mẹo: Đảm bảo rằng file âm thanh của bạn không bị hỏng và API Key của bạn hoạt động bình thường.")
